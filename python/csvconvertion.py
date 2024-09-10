import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def csv_to_docx(csv_file, docx_file):
    # Read the CSV file
    df = pd.read_csv(csv_file)
    
    # Create a new Document
    doc = Document()
    doc.add_heading('CSV Data', 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add table to document
    table = doc.add_table(rows=1, cols=len(df.columns))

    # Add the header row with bold and center alignment
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = column
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add the rows from the dataframe
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
            row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Optional: Adjust cell shading for alternating rows (makes the table more formal)
    for i, row in enumerate(table.rows):
        if i % 2 == 0:  # Shade every other row
            for cell in row.cells:
                cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w'))))

    # Optional: Change font size for the table
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)  # Formal size

    # Save the document
    doc.save(docx_file)
    print(f"CSV data has been converted and saved to {docx_file}")


# Get user input for CSV and DOCX file paths
csv_file_path = input("Enter the path to the CSV file: ")
docx_file_path = input("Enter the path to save the DOCX file (e.g., output.docx): ")

# Convert CSV to DOCX
csv_to_docx(csv_file_path, docx_file_path)
